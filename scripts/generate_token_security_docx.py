from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "ARHIAX_Arquitectura_de_Seguridad_Tokens_Efimeros_y_ADR.docx"

MAIN_MD = DOCS / "ARHIAX_Arquitectura_de_Seguridad_para_Tokens_Efimeros_ES.md"
ADR_MD = DOCS / "ADR-ARHIAX-001-Tokens-Efimeros-y-Delegacion-Gobernada.md"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def ensure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, base_name, size, color, bold in [
        ("ARHIAX Title", "Title", 24, RGBColor(11, 61, 145), True),
        ("ARHIAX Subtitle", "Subtitle", 11, RGBColor(90, 90, 90), False),
        ("ARHIAX Section", "Heading 1", 16, RGBColor(11, 61, 145), True),
        ("ARHIAX Subsection", "Heading 2", 13, RGBColor(26, 26, 26), True),
        ("ARHIAX Note", "Normal", 10, RGBColor(70, 70, 70), False),
    ]:
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles[base_name]
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.space_before = Pt(4)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.style = doc.styles["ARHIAX Note"]
    footer.add_run("ARHIAX Security Architecture ")
    set_page_number(footer)


def add_cover(doc):
    p = doc.add_paragraph(style="ARHIAX Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ARHIAX")
    p.add_run("\nArquitectura de Seguridad para Tokens Efimeros")

    p2 = doc.add_paragraph(style="ARHIAX Subtitle")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(
        "Version ejecutiva detallada en espanol corporativo ARHIAX, con Decision Record "
        "de arquitectura y lineamientos de implementacion."
    )

    doc.add_paragraph("")

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    rows = [
        ("Owner", "Sinergia Consulting Group S.A.S."),
        ("Fecha", "2026-05-13"),
        ("Ambito", "ARHIAX CM, ARHIA-DX, AgentCreator y runtime gobernado"),
        ("Clasificacion", "Arquitectura objetivo de seguridad y gobernanza"),
        ("Estado", "Borrador tecnico detallado para adopcion"),
    ]
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        set_cell_shading(table.cell(i, 0), "D9E8FB")
    for row in table.rows:
        row.cells[0].width = Inches(1.9)
        row.cells[1].width = Inches(4.9)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)

    doc.add_paragraph("")
    note = doc.add_paragraph(style="ARHIAX Note")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        "Este documento consolida la arquitectura objetivo para minimizar el riesgo asociado "
        "a tokens efimeros en entornos multiagente y operaciones de alto impacto."
    )
    doc.add_page_break()


def add_exec_summary(doc):
    doc.add_paragraph("Resumen Ejecutivo", style="ARHIAX Section")
    paras = [
        "ARHIAX debe asumir que un token efimero comprometido conserva valor operacional durante su corta ventana de vida y, por tanto, debe ser tratado como un riesgo critico cuando representa autorizacion sobre actos sensibles, recursos catastrales o ejecuciones de herramientas en agentes.",
        "La postura recomendada no es fortalecer un JWT aislado, sino adoptar una arquitectura de delegacion gobernada en la que cada credencial sea de vida corta, de proposito unico, ligada a posesion cuando sea viable, emitida por un broker, validada en cada frontera, autorizada segun contexto de negocio, aislada del plano LLM y trazable de extremo a extremo.",
        "La pieza central de esta arquitectura es el Credential Broker, que permite desacoplar autenticacion primaria, delegacion, emision efimera y evidencia de uso. Este servicio elimina la necesidad de que el agente o el frontend operen con credenciales amplias o reutilizables.",
        "En ARHIA-DX, la norma no negociable es zero-token-in-context: el modelo nunca debe ver tokens, refresh tokens, API keys, cookies privilegiadas, signed URLs ni headers de autorizacion. El modelo decide intenciones; el plano de credenciales ejecuta.",
    ]
    for text in paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "EEF5FC")
    cell.text = (
        "Decision ejecutiva propuesta: adoptar una arquitectura ARHIAX de delegacion gobernada "
        "para tokens efimeros, con DPoP donde aplique, mTLS interno, revocacion hibrida, "
        "autorizacion contextual y control explicito del plano de herramientas en agentes."
    )
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(4)


def add_markdown(doc, path: Path):
    lines = path.read_text(encoding="utf-8").splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("---"):
            continue
        if line.startswith("```"):
            # skip fence markers; code lines will be handled as plain text
            continue
        if line.startswith("# "):
            doc.add_paragraph(line[2:].strip(), style="ARHIAX Section")
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="ARHIAX Section")
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="ARHIAX Subsection")
            continue
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.add_run(re.match(r"^(\d+\.)", line).group(1) + " ").bold = True
            p.add_run(clean_inline(re.sub(r"^\d+\.\s+", "", line)))
            continue
        if line.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.add_run("• ").bold = True
            p.add_run(clean_inline(line[2:].strip()))
            continue
        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            text = line.strip("*")
            doc.add_paragraph(text, style="ARHIAX Subsection")
            continue
        if line.startswith("```text"):
            continue
        p = doc.add_paragraph()
        if line.startswith("**") and ":" in line:
            p.add_run(clean_inline(line))
        else:
            p.add_run(clean_inline(line))
        p.paragraph_format.space_after = Pt(5)


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("**", "")
    return text


def add_section_break(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_appendix_header(doc, title, subtitle):
    doc.add_paragraph(title, style="ARHIAX Section")
    p = doc.add_paragraph(subtitle, style="ARHIAX Note")
    p.paragraph_format.space_after = Pt(10)


def main():
    doc = Document()
    ensure_styles(doc)
    configure_document(doc)
    add_cover(doc)
    add_exec_summary(doc)
    doc.add_page_break()

    add_appendix_header(
        doc,
        "Documento Base de Arquitectura",
        "Version detallada en espanol corporativo ARHIAX.",
    )
    add_markdown(doc, MAIN_MD)

    doc.add_page_break()
    add_appendix_header(
        doc,
        "Architecture Decision Record",
        "Formalizacion de la decision arquitectonica para adopcion en ARHIAX.",
    )
    add_markdown(doc, ADR_MD)

    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    main()
